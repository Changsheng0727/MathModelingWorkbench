from __future__ import annotations

import csv
import re
from pathlib import Path
from typing import Any, Callable

import pandas as pd
from docx import Document
from openpyxl import load_workbook

from app.services.process_utils import find_external_command, run_external_command


TEXT_EXTENSIONS = {".txt", ".md"}
DOC_EXTENSIONS = {".pdf", ".doc", ".docx", ".txt", ".md"}
DATA_EXTENSIONS = {".xlsx", ".xls", ".csv"}


ProgressCallback = Callable[[dict[str, Any]], None]


def inventory_files(raw_dir: Path, on_file: ProgressCallback | None = None) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    files = [path for path in sorted(raw_dir.rglob("*")) if path.is_file()]
    total = len(files)
    for index, path in enumerate(files, 1):
        suffix = path.suffix.lower()
        rel = path.relative_to(raw_dir).as_posix()
        if on_file:
            on_file({"phase": "start", "index": index, "total": total, "path": rel, "suffix": suffix})
        record: dict[str, Any] = {
            "path": rel,
            "name": path.name,
            "suffix": suffix,
            "size": path.stat().st_size,
            "kind": "document" if suffix in DOC_EXTENSIONS else "data" if suffix in DATA_EXTENSIONS else "other",
        }
        if suffix in DATA_EXTENSIONS:
            record["schema"] = parse_data_schema(path)
        if suffix in DOC_EXTENSIONS:
            text = extract_document_text(path)
            record["text_preview"] = compact_text(text, 700)
            record["char_count"] = len(text)
        records.append(record)
        if on_file:
            on_file(
                {
                    "phase": "finish",
                    "index": index,
                    "total": total,
                    "path": rel,
                    "suffix": suffix,
                    "kind": record["kind"],
                    "size": record["size"],
                }
            )
    return records


def extract_all_document_text(raw_dir: Path, on_document: ProgressCallback | None = None) -> list[dict[str, str]]:
    docs = []
    paths = [path for path in sorted(raw_dir.rglob("*")) if path.is_file() and path.suffix.lower() in DOC_EXTENSIONS]
    total = len(paths)
    for index, path in enumerate(paths, 1):
        rel = path.relative_to(raw_dir).as_posix()
        if on_document:
            on_document({"phase": "start", "index": index, "total": total, "path": rel, "suffix": path.suffix.lower()})
        text = extract_document_text(path)
        docs.append(
            {
                "path": rel,
                "name": path.name,
                "text": text,
            }
        )
        if on_document:
            on_document({"phase": "finish", "index": index, "total": total, "path": rel, "chars": len(text)})
    return docs


def extract_document_text(path: Path) -> str:
    suffix = path.suffix.lower()
    try:
        if suffix == ".pdf":
            return extract_pdf_text(path)
        if suffix == ".docx":
            return extract_docx_text(path)
        if suffix == ".doc":
            return extract_legacy_doc_text(path)
        if suffix in TEXT_EXTENSIONS:
            return read_text_file(path)
    except Exception as exc:
        return f"[解析失败：{type(exc).__name__}: {exc}]"
    return ""


def extract_pdf_text(path: Path) -> str:
    pdftotext = find_external_command("pdftotext")
    if not pdftotext:
        return "[未找到 pdftotext，无法解析 PDF 文本]"
    result = run_external_command(
        [pdftotext, "-layout", "-enc", "UTF-8", str(path), "-"],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        timeout=60,
        check=False,
    )
    return result.stdout if result.stdout else result.stderr


def extract_docx_text(path: Path) -> str:
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_legacy_doc_text(path: Path) -> str:
    pandoc = find_external_command("pandoc")
    if not pandoc:
        return "[未找到 pandoc，无法解析旧版 DOC 文本]"
    result = run_external_command(
        [pandoc, str(path), "-t", "plain"],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        timeout=60,
        check=False,
    )
    if result.returncode == 0 and result.stdout.strip():
        return result.stdout
    fallback = extract_doc_binary_text(path)
    if fallback.strip():
        return fallback
    return result.stdout or result.stderr or "[旧版 DOC 解析失败]"


def extract_doc_binary_text(path: Path) -> str:
    """Best-effort text recovery for legacy binary .doc files.

    The workbench only needs enough readable problem text for analysis when a
    full Office converter is unavailable. Many old Chinese .doc files contain
    recoverable UTF-16LE text runs inside the binary container.
    """
    data = path.read_bytes()
    text = data.decode("utf-16le", errors="ignore")
    spans = re.findall(r"[\u4e00-\u9fffA-Za-z0-9，。；：、（）《》“”！？\s\-_/.%]+", text)
    cleaned = []
    for span in spans:
        span = re.sub(r"\s+", " ", span).strip()
        chinese_count = sum("\u4e00" <= ch <= "\u9fff" for ch in span)
        if chinese_count >= 4 and len(span) >= 12:
            cleaned.append(span)
    return "\n".join(cleaned)


def read_text_file(path: Path) -> str:
    for encoding in ["utf-8-sig", "utf-8", "gbk"]:
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(errors="replace")


def parse_data_schema(path: Path) -> dict[str, Any]:
    suffix = path.suffix.lower()
    if suffix == ".csv":
        return parse_csv_schema(path)
    return parse_excel_schema(path)


def parse_csv_schema(path: Path) -> dict[str, Any]:
    for encoding in ["utf-8-sig", "gbk", "utf-8"]:
        try:
            df = pd.read_csv(path, nrows=5, encoding=encoding)
            rows = count_csv_rows(path, encoding)
            return {
                "type": "csv",
                "encoding": encoding,
                "rows": rows,
                "cols": len(df.columns),
                "columns": [str(c) for c in df.columns],
                "sample": df.head(3).fillna("").astype(str).to_dict(orient="records"),
            }
        except Exception:
            continue
    return {"type": "csv", "error": "无法识别编码或读取失败"}


def count_csv_rows(path: Path, encoding: str) -> int:
    try:
        with path.open("r", encoding=encoding, errors="ignore", newline="") as fh:
            return max(sum(1 for _ in csv.reader(fh)) - 1, 0)
    except Exception:
        return -1


def parse_excel_schema(path: Path) -> dict[str, Any]:
    try:
        workbook = load_workbook(path, read_only=True, data_only=True)
        sheets = []
        for sheet_name in workbook.sheetnames[:12]:
            ws = workbook[sheet_name]
            header = next(ws.iter_rows(min_row=1, max_row=1, values_only=True), ())
            sheets.append(
                {
                    "name": sheet_name,
                    "rows": ws.max_row,
                    "cols": ws.max_column,
                    "columns": [str(x) if x is not None else "" for x in header],
                }
            )
        workbook.close()
        return {"type": "excel", "sheets": sheets}
    except Exception as exc:
        return {"type": "excel", "error": f"{type(exc).__name__}: {exc}"}


def compact_text(text: str, limit: int) -> str:
    text = " ".join(text.split())
    if len(text) <= limit:
        return text
    return text[:limit] + "..."
