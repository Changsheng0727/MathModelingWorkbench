from __future__ import annotations

import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from app.config import APP_ROOT


DOCX_RELATIVE = "paper/main.docx"
WORD_LOG_RELATIVE = "artifacts/word_export.log"
MODULE_PANDOC_REFERENCE_DOC = Path(__file__).resolve().parents[1] / "resources" / "pandoc_reference.docx"
BUNDLED_PANDOC_REFERENCE_DOC = APP_ROOT / "app" / "resources" / "pandoc_reference.docx"
PYINSTALLER_PANDOC_REFERENCE_DOC = APP_ROOT / "_internal" / "app" / "resources" / "pandoc_reference.docx"
EXTERNAL_PANDOC_REFERENCE_DOC = Path(r"E:\AI_MATHMODELING\dongSanShengB\B题\pandoc模板.docx")


def export_word_document(root: Path) -> dict[str, Any]:
    """Export paper/main.tex to paper/main.docx.

    Pandoc is preferred when available because it preserves more structure.
    A python-docx fallback still creates an editable Word document when Pandoc
    is missing or cannot parse the LaTeX source.
    """
    tex_path = root / "paper" / "main.tex"
    if not tex_path.exists():
        raise FileNotFoundError("paper/main.tex 不存在，无法导出 Word")

    docx_path = root / DOCX_RELATIVE
    log_path = root / WORD_LOG_RELATIVE
    log_path.parent.mkdir(parents=True, exist_ok=True)
    logs: list[str] = []

    pandoc = shutil.which("pandoc")
    if pandoc:
        result = run_pandoc(pandoc, tex_path, docx_path, default_pandoc_reference_doc())
        logs.append(result["log"])
        if result["success"]:
            log_path.write_text("\n\n".join(logs), encoding="utf-8")
            return {
                "success": True,
                "method": "pandoc",
                "docx": DOCX_RELATIVE,
                "log": WORD_LOG_RELATIVE,
            }

    fallback = run_python_docx_export(tex_path, docx_path)
    logs.append(fallback["log"])
    log_path.write_text("\n\n".join(logs), encoding="utf-8")
    return {
        "success": fallback["success"],
        "method": "python-docx",
        "docx": DOCX_RELATIVE if fallback["success"] else "",
        "log": WORD_LOG_RELATIVE,
    }


def run_pandoc(pandoc: str, tex_path: Path, docx_path: Path, reference_doc: Path | None = None) -> dict[str, Any]:
    resource_path = ".;..;../results;../artifacts"
    temp_path = write_pandoc_friendly_tex(tex_path)
    command = [pandoc, temp_path.name, "-o", "main.docx", f"--resource-path={resource_path}"]
    reference_note = "未配置 Pandoc reference-doc。"
    if reference_doc:
        if reference_doc.exists() and reference_doc.is_file():
            command.append(f"--reference-doc={reference_doc}")
            reference_note = f"已使用 Pandoc reference-doc：{reference_doc}"
        else:
            reference_note = f"Pandoc reference-doc 不存在，未使用模板：{reference_doc}"
    try:
        result = subprocess.run(
            command,
            cwd=tex_path.parent,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=180,
            check=False,
        )
    finally:
        temp_path.unlink(missing_ok=True)
    success = result.returncode == 0 and docx_path.exists()
    return {
        "success": success,
        "log": "===== pandoc tex to docx =====\n"
        + reference_note
        + "\n"
        + result.stdout
        + "\n"
        + result.stderr
        + ("\nPandoc 导出成功。" if success else "\nPandoc 导出失败，转入 python-docx 兜底导出。"),
    }


def default_pandoc_reference_doc() -> Path:
    """Prefer the reference doc bundled with the desktop app."""
    for candidate in [
        MODULE_PANDOC_REFERENCE_DOC,
        BUNDLED_PANDOC_REFERENCE_DOC,
        PYINSTALLER_PANDOC_REFERENCE_DOC,
        EXTERNAL_PANDOC_REFERENCE_DOC,
    ]:
        if candidate.exists():
            return candidate
    return EXTERNAL_PANDOC_REFERENCE_DOC


def write_pandoc_friendly_tex(tex_path: Path) -> Path:
    tex = tex_path.read_text(encoding="utf-8", errors="replace")
    tex = re.sub(r"\\detokenize\{([^{}]+)\}", r"\1", tex)
    tex = re.sub(r"^\\newcommand\{\\eqnum\}.*$", "", tex, flags=re.M)
    tex = tex.replace(r"\eqnum", "")
    with tempfile.NamedTemporaryFile(
        "w",
        suffix=".tex",
        prefix="word_export_",
        dir=tex_path.parent,
        delete=False,
        encoding="utf-8",
    ) as handle:
        handle.write(tex)
        return Path(handle.name)


def run_python_docx_export(tex_path: Path, docx_path: Path) -> dict[str, Any]:
    tex = tex_path.read_text(encoding="utf-8", errors="replace")
    doc = Document()
    configure_document(doc)
    body = latex_document_body(tex)
    render_body_to_docx(doc, body, tex_path.parent)
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)
    return {
        "success": docx_path.exists(),
        "log": "===== python-docx fallback =====\n已将 LaTeX 正文、标题、公式文本、表格和可识别图片导出为可编辑 Word 文档。",
    }


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.87)
    section.bottom_margin = Inches(0.87)
    section.left_margin = Inches(0.87)
    section.right_margin = Inches(0.87)
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def latex_document_body(tex: str) -> str:
    tex = re.sub(r"(?<!\\)%.*", "", tex)
    match = re.search(r"\\begin\{document\}([\s\S]*?)\\end\{document\}", tex)
    return match.group(1) if match else tex


def render_body_to_docx(doc: Document, body: str, tex_dir: Path) -> None:
    lines = body.splitlines()
    index = 0
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        if not paragraph_buffer:
            return
        text = clean_latex_text(" ".join(paragraph_buffer))
        paragraph_buffer.clear()
        if text:
            doc.add_paragraph(text)

    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if not line:
            flush_paragraph()
            index += 1
            continue

        if line.startswith(r"\begin{table"):
            flush_paragraph()
            block, index = collect_environment(lines, index, "table")
            add_table_block(doc, block)
            continue

        if line.startswith(r"\begin{figure"):
            flush_paragraph()
            block, index = collect_environment(lines, index, "figure")
            add_figure_block(doc, block, tex_dir)
            continue

        if line.startswith("$$"):
            flush_paragraph()
            block, index = collect_display_math(lines, index)
            add_formula(doc, block)
            continue

        heading = parse_heading(line)
        if heading:
            flush_paragraph()
            level, text = heading
            doc.add_heading(clean_latex_text(text), level=level)
            index += 1
            continue

        item = parse_item(line)
        if item:
            flush_paragraph()
            doc.add_paragraph(clean_latex_text(item), style="List Bullet")
            index += 1
            continue

        if should_skip_line(line):
            index += 1
            continue

        paragraph_buffer.append(line)
        index += 1

    flush_paragraph()


def collect_environment(lines: list[str], start: int, name: str) -> tuple[str, int]:
    block = [lines[start]]
    index = start + 1
    end_marker = rf"\end{{{name}}}"
    while index < len(lines):
        block.append(lines[index])
        if end_marker in lines[index]:
            index += 1
            break
        index += 1
    return "\n".join(block), index


def collect_display_math(lines: list[str], start: int) -> tuple[str, int]:
    first = lines[start].strip()
    block = [first[2:]]
    index = start + 1
    if first.endswith("$$") and len(first) > 2:
        return first.strip("$"), index
    while index < len(lines):
        line = lines[index].strip()
        if line.endswith("$$"):
            block.append(line[:-2])
            index += 1
            break
        block.append(line)
        index += 1
    return "\n".join(block), index


def parse_heading(line: str) -> tuple[int, str] | None:
    for command, level in [("section", 1), ("subsection", 2), ("subsubsection", 3)]:
        value = extract_command_argument(line, command)
        if value is not None:
            return level, value
    return None


def parse_item(line: str) -> str | None:
    match = re.match(r"\\item(?:\[[^\]]+\])?\s*(.*)", line)
    return match.group(1) if match else None


def should_skip_line(line: str) -> bool:
    skip_prefixes = [
        r"\thispagestyle",
        r"\pagestyle",
        r"\setcounter",
        r"\phantomsection",
        r"\label",
        r"\newpage",
        r"\clearpage",
        r"\appendix",
        r"\centering",
        r"\begin{center}",
        r"\end{center}",
        r"\begin{enumerate}",
        r"\end{enumerate}",
        r"\begin{itemize}",
        r"\end{itemize}",
    ]
    return any(line.startswith(prefix) for prefix in skip_prefixes)


def add_formula(doc: Document, formula: str) -> None:
    text = clean_latex_text(formula, keep_commands=True)
    if not text:
        return
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")


def add_table_block(doc: Document, block: str) -> None:
    caption = extract_command_argument(block, "caption")
    if caption:
        doc.add_paragraph("表：" + clean_latex_text(caption))
    tabular = re.search(r"\\begin\{tabular\}(?:\[[^\]]*\])?\{[^{}]*\}([\s\S]*?)\\end\{tabular\}", block)
    if not tabular:
        text = clean_latex_text(block)
        if text:
            doc.add_paragraph(text)
        return
    rows = parse_tabular_rows(tabular.group(1))
    if not rows:
        return
    max_cols = min(max(len(row) for row in rows), 8)
    rows = [row[:max_cols] + [""] * max(0, max_cols - len(row)) for row in rows[:80]]
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            table.cell(row_index, col_index).text = clean_latex_text(value)


def parse_tabular_rows(tabular: str) -> list[list[str]]:
    tabular = re.sub(r"\\(?:toprule|midrule|bottomrule|hline|cline\{[^{}]*\})", "", tabular)
    raw_rows = re.split(r"\\\\", tabular)
    rows: list[list[str]] = []
    for raw_row in raw_rows:
        row = raw_row.strip()
        if not row:
            continue
        row = re.sub(r"\\multicolumn\{\d+\}\{[^{}]*\}\{([^{}]*)\}", r"\1", row)
        cells = [cell.strip() for cell in row.split("&")]
        if any(cells):
            rows.append(cells)
    return rows


def add_figure_block(doc: Document, block: str, tex_dir: Path) -> None:
    caption = extract_command_argument(block, "caption")
    include = re.search(r"\\includegraphics(?:\[[^\]]*\])?\{([^{}]+)\}", block)
    image_added = False
    if include:
        image_path = resolve_image(tex_dir, include.group(1))
        if image_path and image_path.suffix.lower() in {".png", ".jpg", ".jpeg", ".bmp"}:
            try:
                doc.add_picture(str(image_path), width=Inches(5.6))
                image_added = True
            except Exception:
                image_added = False
    if caption:
        prefix = "图：" if image_added else "图（原图需在 PDF 或支撑材料中查看）："
        doc.add_paragraph(prefix + clean_latex_text(caption))


def resolve_image(tex_dir: Path, raw_path: str) -> Path | None:
    candidate = (tex_dir / raw_path).resolve()
    if candidate.exists():
        return candidate
    if not candidate.suffix:
        for suffix in [".png", ".jpg", ".jpeg", ".bmp", ".pdf"]:
            with_suffix = candidate.with_suffix(suffix)
            if with_suffix.exists():
                return with_suffix
    return None


def extract_command_argument(text: str, command: str) -> str | None:
    marker = "\\" + command
    start = text.find(marker)
    if start < 0:
        return None
    brace_start = text.find("{", start + len(marker))
    if brace_start < 0:
        return None
    depth = 0
    for index in range(brace_start, len(text)):
        char = text[index]
        if char == "{" and (index == 0 or text[index - 1] != "\\"):
            depth += 1
        elif char == "}" and (index == 0 or text[index - 1] != "\\"):
            depth -= 1
            if depth == 0:
                return text[brace_start + 1 : index]
    return None


def clean_latex_text(text: str, keep_commands: bool = False) -> str:
    text = text.replace("\n", " ")
    text = re.sub(r"\\label\{[^{}]*\}", "", text)
    text = re.sub(r"\\ref\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\cite\{([^{}]*)\}", r"[\1]", text)
    for command in ["textbf", "emph", "textit", "heiti", "songti", "kaishu"]:
        text = replace_one_arg_command(text, command)
    text = re.sub(r"\\zihao\{[^{}]*\}", "", text)
    text = re.sub(r"\\href\{[^{}]*\}\{([^{}]*)\}", r"\1", text)
    text = text.replace(r"\%", "%").replace(r"\&", "&").replace(r"\_", "_")
    text = text.replace(r"\#", "#").replace(r"\$", "$").replace(r"\{", "{").replace(r"\}", "}")
    replacements = {
        r"\times": "×",
        r"\cdot": "·",
        r"\leq": "≤",
        r"\geq": "≥",
        r"\alpha": "alpha",
        r"\beta": "beta",
        r"\gamma": "gamma",
        r"\theta": "theta",
        r"\lambda": "lambda",
        r"\mu": "mu",
        r"\sigma": "sigma",
        r"\Omega": "Omega",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = text.replace("$$", "").replace(r"\[", "").replace(r"\]", "")
    text = re.sub(r"\$([^$]+)\$", r"\1", text)
    if not keep_commands:
        text = re.sub(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?(?:\{([^{}]*)\})", r"\1", text)
        text = re.sub(r"\\[a-zA-Z]+\*?", "", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def replace_one_arg_command(text: str, command: str) -> str:
    pattern = re.compile(rf"\\{command}\{{([^{{}}]*)\}}")
    previous = None
    while previous != text:
        previous = text
        text = pattern.sub(r"\1", text)
    return text
